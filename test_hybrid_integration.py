"""集成验证：text_normalizer + parser + chunking + BM25 + hybrid retrieval 全链路。

mock 掉 streamlit / zhipuai / dotenv / chromadb 等外部依赖，
但保留 parser -> normalizer -> chunking -> document_router -> bm25 -> retrieval 的真实调用。
"""
import os
import sys
import tempfile
from unittest.mock import MagicMock, patch, PropertyMock

# ── mock 外部重依赖 ──────────────────────────────────────────────
mock_st = MagicMock()
mock_st.cache_resource = lambda f: f
sys.modules.setdefault("streamlit", mock_st)

mock_zhipuai = MagicMock()
sys.modules.setdefault("zhipuai", mock_zhipuai)

mock_dotenv = MagicMock()
mock_dotenv.load_dotenv = lambda: None
sys.modules.setdefault("dotenv", mock_dotenv)

mock_chromadb = MagicMock()
sys.modules.setdefault("chromadb", mock_chromadb)

# ── 导入业务模块 ─────────────────────────────────────────────────
from services.text_normalizer import normalize_pdf_text, normalize_plain_text
from services.parser import extract_document_pages
from services.chunking import build_page_chunks
from services.document_router import detect_doc_category, choose_chunk_strategy
from services.bm25_index import rebuild_bm25_index, bm25_search, is_index_ready
from services.retrieval import rrf_fuse
import services.rag_pipeline as pipeline


# ── 构造临时文件 ─────────────────────────────────────────────────

def make_txt(content="Payment terms: Net 30 days.\n\nThe vendor shall deliver within 5 business days."):
    f = tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False, encoding="utf-8")
    f.write(content)
    f.close()
    return f.name


def make_docx():
    from docx import Document
    f = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    path = f.name
    f.close()
    doc = Document()
    doc.add_paragraph("Section 1: Scope of Work")
    doc.add_paragraph("The contractor shall provide all labor and materials.")
    doc.save(path)
    return path


def make_pdf():
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(w=0, h=10, text="Article 1: General Terms\nThe lessor agrees to provide access.\nPayment shall be made monthly.")
    f = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    path = f.name
    f.close()
    pdf.output(path)
    return path


# ── 测试函数 ─────────────────────────────────────────────────────

def test_normalizer():
    raw = "This is a test  \t line.\r\n\r\n\r\nAnother   paragraph."
    result = normalize_plain_text(raw)
    assert "\r" not in result
    assert "\t" not in result
    assert "\n\n\n" not in result
    print("✓ text_normalizer: plain text 清洗正常")

    pdf_raw = "header line\nSome contract text that\ncontinues on next line.\n\nNew paragraph here."
    result = normalize_pdf_text(pdf_raw)
    assert result.strip()
    print("✓ text_normalizer: PDF text 标准化正常")


def test_parser_normalized_text():
    for factory, ext in [(make_txt, "txt"), (make_docx, "docx"), (make_pdf, "pdf")]:
        path = factory()
        try:
            pages = extract_document_pages(path)
            assert pages, f"{ext}: 解析结果为空"
            p = pages[0]
            assert "text" in p, f"{ext}: 缺少 text"
            assert "normalized_text" in p, f"{ext}: 缺少 normalized_text"
            assert "file_type" in p, f"{ext}: 缺少 file_type"
            assert p["file_type"] == ext, f"{ext}: file_type={p['file_type']}"
            assert p["normalized_text"].strip(), f"{ext}: normalized_text 为空"
            print(f"✓ parser: {ext.upper()} 返回含 normalized_text + file_type")
        finally:
            os.unlink(path)


def test_chunking_uses_normalized():
    path = make_txt("Short text for chunking test.\n\nAnother paragraph with more words to make it work.")
    try:
        pages = extract_document_pages(path)
        page = pages[0]
        page["source"] = "test.txt"
        chunks = build_page_chunks(page, chunk_strategy="paragraph")
        assert chunks, "chunking: 无 chunk 产出"
        assert chunks[0]["text"].strip()
        print(f"✓ chunking: 基于 normalized_text 切块正常，产出 {len(chunks)} 个 chunks")
    finally:
        os.unlink(path)


def test_bm25_and_rrf():
    docs = [
        "Payment terms are net 30 days from invoice date",
        "The contractor shall deliver materials on time",
        "Force majeure clause applies to natural disasters",
    ]
    metas = [
        {"source": "a.pdf", "page": 1, "chunk_index": 0, "file_type": "pdf",
         "doc_category": "natural_pdf", "chunk_strategy": "paragraph"},
        {"source": "a.pdf", "page": 1, "chunk_index": 1, "file_type": "pdf",
         "doc_category": "natural_pdf", "chunk_strategy": "paragraph"},
        {"source": "b.txt", "page": 1, "chunk_index": 0, "file_type": "txt",
         "doc_category": "natural_text", "chunk_strategy": "paragraph"},
    ]

    rebuild_bm25_index(docs, metas)
    assert is_index_ready(), "BM25 索引未就绪"

    results = bm25_search("payment terms", top_k=2)
    assert results, "BM25 搜索无结果"
    assert results[0]["bm25_score"] > 0
    assert results[0]["source"] == "a.pdf"
    print(f"✓ BM25: 搜索 'payment terms' 返回 {len(results)} 条，top1 score={results[0]['bm25_score']:.4f}")

    # 模拟 dense 结果
    dense_results = [
        {"text": docs[2], "source": "b.txt", "page": 1, "chunk_index": 0,
         "file_type": "txt", "doc_category": "natural_text", "chunk_strategy": "paragraph",
         "dense_score": 0.5, "dense_rank": 1},
        {"text": docs[0], "source": "a.pdf", "page": 1, "chunk_index": 0,
         "file_type": "pdf", "doc_category": "natural_pdf", "chunk_strategy": "paragraph",
         "dense_score": 0.7, "dense_rank": 2},
    ]

    fused = rrf_fuse(dense_results, results, top_k=3)
    assert fused, "RRF 融合无结果"
    for item in fused:
        assert "fused_rank" in item
        assert "rrf_score" in item
        assert "dense_score" in item
        assert "bm25_score" in item
    print(f"✓ RRF: 融合产出 {len(fused)} 条，top1 rrf_score={fused[0]['rrf_score']:.6f}")


def test_pipeline_index():
    """验证 index_document 完整链路（mock vector store）。"""
    path = make_txt("Delivery shall be completed within 14 days.\n\nPenalty for delay is 1% per week.")
    captured = {}

    from services.document_router import detect_doc_category as real_detect

    def spy_detect(file_type, pages):
        captured["file_type"] = file_type
        captured["doc_category"] = real_detect(file_type=file_type, pages=pages)
        return captured["doc_category"]

    mock_collection = MagicMock()

    try:
        with patch.object(pipeline, "init_collection", return_value=mock_collection), \
             patch.object(pipeline, "delete_document_by_source"), \
             patch.object(pipeline, "detect_doc_category", side_effect=spy_detect), \
             patch.object(pipeline, "get_all_chunks", return_value=([], [])):

            count = pipeline.index_document(path)

        assert captured["file_type"] == "txt"
        assert count > 0
        assert mock_collection.add.called

        metadatas = mock_collection.add.call_args.kwargs.get("metadatas", [])
        for m in metadatas:
            assert m["file_type"] == "txt"

        print(f"✓ pipeline: index_document 链路正常，file_type='txt', chunks={count}")
    finally:
        os.unlink(path)


if __name__ == "__main__":
    test_normalizer()
    test_parser_normalized_text()
    test_chunking_uses_normalized()
    test_bm25_and_rrf()
    test_pipeline_index()
    print("\n=============================")
    print("全部集成验证通过 ✓")
    print("=============================")
