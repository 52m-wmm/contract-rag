"""集成验证：rag_pipeline.index_document -> document_router 链路中 file_type 是否正确传递。

mock 掉 vector store / embedding / streamlit（避免外部依赖），
但保留 parser -> pipeline -> document_router -> chunking 的真实调用。
"""
import os
import sys
import tempfile
from unittest.mock import MagicMock, patch

# ── 在导入业务模块前，先 mock 掉重外部依赖 ──────────────────────
# streamlit（vector_store 用了 @st.cache_resource）
mock_st = MagicMock()
mock_st.cache_resource = lambda f: f  # 装饰器透传
sys.modules.setdefault("streamlit", mock_st)

# zhipuai（embeddings + rag_pipeline 顶层初始化）
mock_zhipuai_module = MagicMock()
sys.modules.setdefault("zhipuai", mock_zhipuai_module)

# dotenv
mock_dotenv = MagicMock()
mock_dotenv.load_dotenv = lambda: None
sys.modules.setdefault("dotenv", mock_dotenv)

# chromadb（vector_store + embeddings）
mock_chromadb = MagicMock()
sys.modules.setdefault("chromadb", mock_chromadb)

# 现在可以安全导入业务模块
from services.document_router import detect_doc_category as real_detect
import services.rag_pipeline as rag_pipeline_mod


# ── 构造临时测试文件 ──────────────────────────────────────────────

def make_txt():
    f = tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False, encoding="utf-8")
    f.write("This is a test text file.\n\nSecond paragraph for chunking.")
    f.close()
    return f.name


def make_docx():
    from docx import Document
    f = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    path = f.name
    f.close()
    doc = Document()
    doc.add_paragraph("This is a test docx file.")
    doc.add_paragraph("Second paragraph for chunking.")
    doc.save(path)
    return path


def make_pdf():
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(text="This is a test PDF file.")
    f = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    path = f.name
    f.close()
    pdf.output(path)
    return path


# ── 测试主体 ─────────────────────────────────────────────────────

def run_test(file_factory, expected_type):
    path = file_factory()
    captured = {}

    # spy：捕获 file_type 后调用真实实现
    def spy_detect(file_type, pages):
        captured["file_type"] = file_type
        return real_detect(file_type=file_type, pages=pages)

    mock_collection = MagicMock()

    try:
        with patch.object(rag_pipeline_mod, "init_collection", return_value=mock_collection), \
             patch.object(rag_pipeline_mod, "delete_document_by_source"), \
             patch.object(rag_pipeline_mod, "detect_doc_category", side_effect=spy_detect):

            count = rag_pipeline_mod.index_document(path)

        ft = captured.get("file_type")
        assert ft == expected_type, f"期望 '{expected_type}'，实际 '{ft}'"

        # 额外验证：写入 vector store 的 metadata 里 file_type 也正确
        if mock_collection.add.called:
            call_kwargs = mock_collection.add.call_args
            metadatas = call_kwargs.kwargs.get("metadatas") or call_kwargs[1].get("metadatas")
            if metadatas:
                for m in metadatas:
                    assert m["file_type"] == expected_type, \
                        f"metadata file_type 期望 '{expected_type}'，实际 '{m['file_type']}'"

        print(f"✓ {expected_type.upper():4s} -> detect_doc_category 收到 file_type='{ft}', chunks={count}")

    finally:
        os.unlink(path)


if __name__ == "__main__":
    run_test(make_txt,  "txt")
    run_test(make_docx, "docx")
    run_test(make_pdf,  "pdf")
    print("\n全部集成验证通过 ✓")
